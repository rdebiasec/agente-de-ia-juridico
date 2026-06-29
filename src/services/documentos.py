"""Servicio de documentos: generación de .docx/.pdf y extracción de texto.

El .docx usa python-docx (puro Python). El .pdf usa WeasyPrint (requiere libs de
sistema: pango/cairo; presentes en el Dockerfile y, en Mac, vía Homebrew).
"""

from __future__ import annotations

import html
from io import BytesIO

from src.storage.models import Draft


def generar_docx(titulo: str, secciones: list[tuple[str, str]] | None = None, cuerpo: str = "") -> bytes:
    """Genera un .docx en memoria. `secciones` = [(encabezado, texto), ...]."""
    from docx import Document

    doc = Document()
    if titulo:
        doc.add_heading(titulo, level=0)

    if cuerpo:
        for parrafo in cuerpo.split("\n\n"):
            doc.add_paragraph(parrafo.strip())

    for encabezado, texto in secciones or []:
        if encabezado:
            doc.add_heading(encabezado, level=1)
        for parrafo in (texto or "").split("\n\n"):
            doc.add_paragraph(parrafo.strip())

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def docx_desde_borrador(draft: Draft) -> bytes:
    """Materializa un borrador aprobado/editado como documento .docx."""
    return generar_docx(titulo=draft.titulo or draft.tipo, cuerpo=draft.contenido)


_PDF_DISCLAIMER = (
    "Borrador informativo — requiere revisión y aprobación del abogado del despacho."
)


def _html_borrador(draft: Draft) -> str:
    titulo = html.escape(draft.titulo or draft.tipo)
    parrafos = "".join(
        f"<p>{html.escape(p.strip())}</p>" for p in draft.contenido.split("\n\n") if p.strip()
    )
    meta = html.escape(f"Tipo: {draft.tipo}" + (f" · Materia: {draft.materia}" if draft.materia else ""))
    return f"""<!doctype html>
<html lang="es"><head><meta charset="utf-8">
<style>
  @page {{ size: Letter; margin: 2.5cm; }}
  body {{ font-family: 'DejaVu Serif', serif; font-size: 12pt; color: #111; line-height: 1.5; }}
  h1 {{ font-size: 16pt; border-bottom: 2px solid #333; padding-bottom: 6px; }}
  .meta {{ color: #555; font-size: 10pt; margin-bottom: 18px; }}
  .disclaimer {{ margin-top: 30px; padding-top: 10px; border-top: 1px solid #999;
    color: #555; font-size: 9pt; font-style: italic; }}
</style></head>
<body>
  <h1>{titulo}</h1>
  <div class="meta">{meta}</div>
  {parrafos}
  <div class="disclaimer">{html.escape(_PDF_DISCLAIMER)}</div>
</body></html>"""


def generar_pdf_desde_borrador(draft: Draft) -> bytes:
    """Genera un PDF del borrador con membrete y disclaimer (WeasyPrint)."""
    from weasyprint import HTML

    return HTML(string=_html_borrador(draft)).write_pdf()


def extraer_texto(nombre: str, data: bytes) -> str:
    """Extrae texto de un adjunto (.pdf, .docx, .txt) para análisis/RAG."""
    lower = (nombre or "").lower()
    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if lower.endswith(".docx"):
        from docx import Document

        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="replace").strip()
    raise ValueError(f"Formato no soportado para extracción: {nombre}")
